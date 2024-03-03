import streamlit as st
import os
import docx
import zipfile
from elevenlabs import set_api_key, voices, generate, save, Voice, VoiceSettings
import base64
import re
import shutil


# Diretório de saída para os arquivos de narração
diretorio_audio = "audio_narracao"

# Substituições de palavras desejadas
substituicoes = [
    ("Hely", "Elí"),
    ("Di Pietro", "Di piêtro"),
    ("CF", "Constituição Federal"),
    ("nº", "número"),
    ("art.", "artigo"),
    ("obs.", "observação"),
    ("J.J.", "José Joaquim"),
    # Adicione outras substituições conforme necessário
]

# Inicializar st.session_state.stop_execution
if "stop_execution" not in st.session_state:
    st.session_state.stop_execution = False

# Função para aplicar substituições de palavras
def aplicar_substituicoes(texto, substituicoes):
    for palavra_antiga, palavra_nova in substituicoes:
        texto = texto.replace(palavra_antiga, palavra_nova)
    return texto

# Crie uma pasta para os arquivos de narração se ela não existir
if not os.path.exists(diretorio_audio):
    os.mkdir(diretorio_audio)

def generate_audio_filename(title, narrator_number, role):
    if title:
        return f"{narrator_number:02d}.{role}_{title}.mp3"
    else:
        return f"{narrator_number:02d}.{role}.mp3"


def generate_audio(texto, voice_id, api_key):
    set_api_key(api_key)
    try:
        audio = generate(
            text=texto,
            voice=Voice(
                voice_id=voice_id,
                settings=VoiceSettings(stability=1.0, similarity_boost=0.70, style=0.0, use_speaker_boost=True)
            ),
            model="eleven_multilingual_v2"
        )
        return audio
    except Exception as e:
        st.warning(f"Ocorreu um problema ao gerar áudio: {str(e)}")
        return None

def save_audio(texto, filename, voice_id, api_key):
    try:
        audio = generate_audio(texto, voice_id, api_key)

        if audio is not None:
            save(
                audio=audio,
                filename=filename
            )
            st.audio(filename, format="audio/mp3")
        else:
            st.error("Desculpe, não há créditos suficientes para gerar áudio. Por favor, recarregue seus créditos.")

    except :
        # Lidar com outras exceções, se necessário
        st.error(f"Ocorreu um erro ao gerar áudio:")


def print_last_voice_info(api_key):
    set_api_key(api_key)
    all_voices = voices()

    st.write("Última Voz Disponível:")
    last_voice = all_voices[-1] if all_voices else None
    st.write(last_voice)

def processar_narrador(texto, narrator_number, voice_id_professor, api_key_professor, voice_id_aluno, api_key_aluno):
    # Identificar o narrador usando expressões regulares
    match_professor = re.search(r'\*\*Professor Van Scheffelt:?\*\*|\*Professor Van Scheffelt\*|\*\*Professor Van Scheffelt\*\*|\*\*professor van scheffelt:?\*\*|\*professor van scheffelt\*|\*\*professor van scheffelt\*\*|Professor Van Scheffelt:', texto, re.IGNORECASE)
    match_aluno = re.search(r'\*\*Aluno:?\*\*|\*Aluno\*|\*\*Aluno\*\*|\*\*aluno:?\*\*|\*aluno\*|\*\*aluno\*\*|Aluno:', texto, re.IGNORECASE)

    if match_professor:
        parts = re.split(r'\*\*Professor Van Scheffelt:?\*\*|\*Professor Van Scheffelt\*|\*\*Professor Van Scheffelt\*\*|\*\*professor van scheffelt:?\*\*|\*professor van scheffelt\*|\*\*professor van scheffelt\*\*|Professor Van Scheffelt:', texto, re.IGNORECASE)
        role = "Professor"
        voice_id = voice_id_professor
        api_key = api_key_professor
    elif match_aluno:
        parts = re.split(r'\*\*Aluno:?\*\*|\*Aluno\*|\*\*Aluno\*\*|\*\*aluno:?\*\*|\*aluno\*|\*\*aluno\*\*|Aluno:', texto, re.IGNORECASE)
        role = "Aluno"
        voice_id = voice_id_aluno
        api_key = api_key_aluno
    else:
        st.warning(f"Não foi possível identificar o narrador para o texto: {texto}. Pulando para o próximo parágrafo.")
        return

    # Restante da função permanece inalterado
    processar_parte_do_texto(parts, narrator_number, role, voice_id, api_key)

# Função para processar uma parte específica do texto
def processar_parte_do_texto(parts, narrator_number, role, voice_id, api_key):
    # Verifica se há pelo menos dois elementos em parts
    if len(parts) >= 2:
        title = parts[0].strip()
        title = title.replace('?', '').replace(':', '').replace('*', '')

        audio_filename = generate_audio_filename(title, narrator_number, role)
        audio_filename = os.path.join(diretorio_audio, audio_filename)

        if os.path.exists(audio_filename):
            st.write(f"O arquivo {audio_filename} já existe na pasta. Pular etapa.")
        else:
            pass

        texto_narracao = aplicar_substituicoes(parts[1], substituicoes)
        save_audio(texto_narracao, audio_filename, voice_id, api_key)
        st.audio(audio_filename, format="audio/mp3")

        # Mostra o texto criado
        st.success(f"Texto criado para {role} {narrator_number}:\n\n{texto_narracao}")
    else:
        st.warning(f"Não foi possível processar o texto: {texto_narracao}. Não há conteúdo suficiente.")


def get_binary_file_downloader_html(filename):
    with open(filename, 'rb') as f:
        zip_data = f.read()

    zip_base64 = base64.b64encode(zip_data).decode()
    href = f'<a href="data:application/zip;base64,{zip_base64}" download="{filename}">Baixar Arquivos ZIP</a>'
    return href

def main():
    st.title("Aplicativo de Narração")

    # Campos para inserir chaves API e voice IDs
    api_key_professor = st.text_input("Digite a chave API do Professor:")
    voice_id_professor = st.text_input("Digite o voice_id do Professor:")
    
    api_key_aluno = st.text_input("Digite a chave API do Aluno:")
    voice_id_aluno = st.text_input("Digite o voice_id do Aluno:")

    # Criação de abas
    with st.expander("Ver voice_id"):
        st.header("Informações da Voz do Professor")
        mostrar_info_professor = st.button("Mostrar Informações da Voz do Professor")
        if mostrar_info_professor:
            print_last_voice_info(api_key_professor)
        else:
            st.warning("Por favor, insira uma chave API válida.")
        st.header("Informações da Voz do Aluno")
        mostrar_info_aluno = st.button("Mostrar Informações da Voz do Aluno")
        if mostrar_info_aluno:
            print_last_voice_info(api_key_aluno)
        else:
            st.warning("Por favor, insira uma chave API válida.")

    # Abra o arquivo do Word
    arquivo_word = st.file_uploader("Selecione o arquivo Word (.docx):", type=["docx"])
    
    st.header("Visualizar Texto do Arquivo")

    if arquivo_word is not None:
        doc = docx.Document(arquivo_word)
        texto_arquivo = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        st.text_area("Texto do Arquivo", texto_arquivo, height=300)

    # Botão para iniciar o processo de criação das narrações
    iniciar_criacao_button = st.button("Iniciar Criação de Narrações")

    # Botão para parar a execução
    parar_execucao_button = st.button("Parar Execução")

    # Adicionar botão para limpar a pasta
    limpar_pasta_button = st.button("Limpar Pasta de Áudios")
    if limpar_pasta_button:
        try:
            shutil.rmtree(diretorio_audio)
            os.makedirs(diretorio_audio)
            st.success("Pasta de áudios limpa com sucesso.")
        except Exception as e:
            st.warning(f"Erro ao limpar pasta de áudios: {e}")

    if iniciar_criacao_button and arquivo_word:
        if api_key_professor and voice_id_professor and api_key_aluno and voice_id_aluno:
            st.write("Iniciando o processo de criação de narrações...")

            doc = docx.Document(arquivo_word)
            paragrafos = []
            narrator_number = 1

            for paragraph in doc.paragraphs:
                if st.session_state.stop_execution:
                    st.warning("Execução interrompida.")
                    st.session_state.stop_execution = False
                    break

                texto = paragraph.text
                if texto.strip():
                    paragrafos.append(texto)

                    processar_narrador(texto, narrator_number, voice_id_professor, api_key_professor, voice_id_aluno, api_key_aluno)
                    narrator_number += 1

            st.write("Processo de criação de narrações concluído!")

            # Criar arquivo ZIP com os arquivos de áudio
            zip_filename = "audio_narracao.zip"
            with zipfile.ZipFile(zip_filename, 'w') as zipf:
                for root, dirs, files in os.walk(diretorio_audio):
                    for file in files:
                        file_path = os.path.join(root, file)
                        zipf.write(file_path, os.path.relpath(file_path, diretorio_audio))

            # Adicionar botão para download do arquivo ZIP
            st.markdown(get_binary_file_downloader_html(zip_filename), unsafe_allow_html=True)

            

if __name__ == "__main__":
    main()
